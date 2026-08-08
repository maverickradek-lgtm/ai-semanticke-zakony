"""
Sync FS pokyny (Pokyny D + Pokyny MF, financnisprava.gov.cz) do samostatne
Neon databaze. Stejny vzor jako scripts/sync_mf_metodiky.py (Financni sprava
je dalsi zdroj "mimo e-Sbirku" s vlastnim Neon projektem zdarma).

Rozdil oproti MF: kazda rocni stranka uz primo obsahuje nazev, datum, odkaz
na soubor a popis (vcetne pripadne poznamky o zruseni/nahrazeni v <strong>
tagu) - zadna zvlastni detailni stranka se nenacita, staci jeden pruchod
pres kazdou rocni stranku obou serii.

Vsechny retezce v tomto souboru jsou schvalne bez ceskych znaku (pouzito
chr() tam, kde je diakritika potreba za behu) kvuli mojibake bugu pri
prenosu pres GitHub Contents API (viz projektova pamet).
"""

import os
import time
from datetime import datetime, date
from io import BytesIO

import requests
from bs4 import BeautifulSoup
from pypdf import PdfReader
from docx import Document
import psycopg2

NL = chr(10)
TAB = chr(9)
SCARON = chr(0x161)

NEON_DB_URL = os.environ["NEON_FS_DB_URL"]
SUPABASE_URL = os.environ["SUPABASE_URL"].rstrip("/")
SERVICE_KEY = os.environ["SUPABASE_SERVICE_ROLE_KEY"]
ADMIN_USER_ID = "2648f5db-bea6-4cac-b490-ad0ec59723df"
GEMINI_API_KEY_OVERRIDE = os.environ.get("GEMINI_API_KEY_OVERRIDE")

EMBED_MODEL = "gemini-embedding-001"
EMBED_DIM = 256
MAX_CHARS_PER_CHUNK = 3000
MAX_NEW_DOCS_PER_RUN = int(os.environ.get("MAX_NEW_DOCS_PER_RUN", "40"))
TIME_BUDGET_SECONDS = int(os.environ.get("TIME_BUDGET_SECONDS", "3000"))
START_YEAR = int(os.environ.get("START_YEAR", "1993"))

BASE_URL = "https://financnisprava.gov.cz"
SERIES_CONFIG = {
    "pokyny_d": BASE_URL + "/cs/dane/legislativa-a-metodika/pokyny-d/casove-cleneni/{year}",
    "pokyny_mf": BASE_URL + "/cs/dane/legislativa-a-metodika/pokyny-mf/{year}",
}
REQ_HEADERS = {"User-Agent": "Mozilla/5.0 (compatible; AIsemantickeZakony/1.0)"}

START_TIME = time.time()


def time_left():
    return TIME_BUDGET_SECONDS - (time.time() - START_TIME)


def sb_headers():
    return {
        "apikey": SERVICE_KEY,
        "Authorization": "Bearer " + SERVICE_KEY,
        "Content-Type": "application/json",
    }


def get_admin_gemini_key():
    if GEMINI_API_KEY_OVERRIDE:
        return GEMINI_API_KEY_OVERRIDE
    r = requests.post(
        SUPABASE_URL + "/rest/v1/rpc/get_user_gemini_key",
        headers=sb_headers(),
        json={"p_user_id": ADMIN_USER_ID},
        timeout=30,
    )
    r.raise_for_status()
    key = r.json()
    if not key:
        raise RuntimeError("Admin Gemini key not available")
    return key


def ensure_schema(conn):
    with conn.cursor() as cur:
        cur.execute("create extension if not exists pgcrypto")
        cur.execute("create extension if not exists vector")
        cur.execute(
            """
            create table if not exists documents (
                id uuid primary key default gen_random_uuid(),
                source text not null default 'fs',
                series text,
                external_id text not null,
                title text not null,
                url text not null,
                pdf_url text,
                published_date date,
                is_current boolean not null default true,
                fetched_at timestamptz default now(),
                created_at timestamptz default now(),
                unique(source, external_id)
            )
            """
        )
        cur.execute(
            """
            create table if not exists chunks (
                id uuid primary key default gen_random_uuid(),
                document_id uuid not null references documents(id) on delete cascade,
                chunk_index int not null,
                heading text,
                content text,
                embedding vector(256),
                created_at timestamptz default now()
            )
            """
        )
        cur.execute("create index if not exists chunks_document_id_idx on chunks(document_id)")
    conn.commit()


def fetch_existing_external_ids(conn, source="fs"):
    with conn.cursor() as cur:
        cur.execute("select external_id from documents where source = %s", (source,))
        return set(row[0] for row in cur.fetchall())


def parse_date(text):
    if not text:
        return None
    parts = [p.strip() for p in text.strip().split(".")]
    nums = []
    for p in parts:
        if p.isdigit():
            nums.append(int(p))
    if len(nums) < 3:
        return None
    day, month, year = nums[0], nums[1], nums[2]
    try:
        return date(year, month, day)
    except ValueError:
        return None


def is_superseded(strong_text):
    s = strong_text.lower().replace(SCARON, "s")
    return ("zrusen" in s) or ("nahrazen" in s)


def list_year(series, year):
    url = SERIES_CONFIG[series].format(year=year)
    resp = requests.get(url, headers=REQ_HEADERS, timeout=30)
    if resp.status_code == 404:
        return []
    resp.raise_for_status()
    soup = BeautifulSoup(resp.text, "html.parser")
    items = []
    for teaser in soup.select(".b-teaser"):
        title_el = teaser.select_one(".b-teaser__title")
        if not title_el:
            continue
        title = title_el.get_text(strip=True)
        link_el = teaser.select_one(".b-teaser__extra a[href]")
        if not link_el:
            continue
        href = link_el.get("href")
        if not href:
            continue
        file_url = href if href.startswith("http") else BASE_URL + href
        fname = file_url.rsplit("/", 1)[-1]
        ext = fname.rsplit(".", 1)[-1].lower() if "." in fname else ""
        if ext not in ("pdf", "docx"):
            continue
        meta_el = teaser.select_one(".b-teaser__meta")
        published = parse_date(meta_el.get_text(strip=True)) if meta_el else None
        text_el = teaser.select_one(".b-teaser__text")
        description = text_el.get_text(" ", strip=True) if text_el else ""
        strong_text = " ".join(s.get_text(" ", strip=True) for s in (text_el.select("strong") if text_el else []))
        current_flag = not is_superseded(strong_text)
        stem = fname.rsplit(".", 1)[0]
        external_id = series + ":" + stem.lower()
        items.append({
            "external_id": external_id,
            "title": title,
            "url": url,
            "file_url": file_url,
            "file_kind": ext,
            "published_date": published,
            "is_current": current_flag,
            "description": description,
        })
    return items


def extract_pdf_text(pdf_bytes):
    reader = PdfReader(BytesIO(pdf_bytes))
    parts = []
    for page in reader.pages:
        t = page.extract_text() or ""
        parts.append(t)
    text = NL.join(parts)
    text = text.replace(TAB, " ")
    while "   " in text:
        text = text.replace("   ", "  ")
    while (NL + NL + NL) in text:
        text = text.replace(NL + NL + NL, NL + NL)
    return text.strip()


def extract_docx_text(docx_bytes):
    doc = Document(BytesIO(docx_bytes))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    parts.append(cell.text.strip())
    return NL.join(parts).strip()


def split_into_chunks(text, max_chars=MAX_CHARS_PER_CHUNK):
    sep = NL + NL
    paragraphs = [p.strip() for p in text.split(sep) if p.strip()]
    chunks = []
    current = []
    current_len = 0
    for p in paragraphs:
        if current_len + len(p) > max_chars and current:
            chunks.append(sep.join(current))
            current = []
            current_len = 0
        current.append(p)
        current_len += len(p)
    if current:
        chunks.append(sep.join(current))
    if not chunks and text.strip():
        chunks = [text[i:i + max_chars] for i in range(0, len(text), max_chars)]
    return chunks


def embed_text(text, gemini_key, retries=3):
    for attempt in range(retries):
        try:
            resp = requests.post(
                "https://generativelanguage.googleapis.com/v1beta/models/" + EMBED_MODEL + ":embedContent",
                headers={"Content-Type": "application/json", "x-goog-api-key": gemini_key},
                json={
                    "content": {"parts": [{"text": text[:8000]}]},
                    "taskType": "RETRIEVAL_DOCUMENT",
                    "outputDimensionality": EMBED_DIM,
                },
                timeout=30,
            )
            if resp.status_code == 429:
                time.sleep(5 * (attempt + 1))
                continue
            resp.raise_for_status()
            vec = resp.json().get("embedding", {}).get("values")
            if not vec:
                return None
            return vec
        except requests.RequestException:
            if attempt == retries - 1:
                raise
            time.sleep(3 * (attempt + 1))
    return None


def import_new_documents(conn):
    existing = fetch_existing_external_ids(conn)
    current_year = datetime.now().year
    new_count = 0
    with conn.cursor() as cur:
        for series in SERIES_CONFIG:
            for year in range(current_year, START_YEAR - 1, -1):
                if time_left() <= 0 or new_count >= MAX_NEW_DOCS_PER_RUN:
                    break
                try:
                    items = list_year(series, year)
                except requests.RequestException as e:
                    print("WARN: list_year failed: " + series + " " + str(year) + " " + str(e))
                    continue
                for item in items:
                    if time_left() <= 0 or new_count >= MAX_NEW_DOCS_PER_RUN:
                        break
                    if item["external_id"] in existing:
                        continue
                    try:
                        file_resp = requests.get(item["file_url"], headers=REQ_HEADERS, timeout=60)
                        file_resp.raise_for_status()
                        if item["file_kind"] == "docx":
                            body_text = extract_docx_text(file_resp.content)
                        else:
                            body_text = extract_pdf_text(file_resp.content)
                    except Exception as e:
                        print("WARN: download/extract failed: " + item["title"] + " " + str(e))
                        existing.add(item["external_id"])
                        continue
                    full_text = item["description"] + NL + NL + body_text if item["description"] else body_text
                    if not full_text or len(full_text) < 20:
                        print("SKIP (no text): " + item["title"])
                        existing.add(item["external_id"])
                        continue
                    cur.execute(
                        "insert into documents (source, series, external_id, title, url, pdf_url, published_date, is_current) values (%s,%s,%s,%s,%s,%s,%s,%s) returning id",
                        ("fs", series, item["external_id"], item["title"], item["url"], item["file_url"], item["published_date"], item["is_current"]),
                    )
                    doc_id = cur.fetchone()[0]
                    chunk_list = split_into_chunks(full_text)
                    for i, ch in enumerate(chunk_list):
                        cur.execute(
                            "insert into chunks (document_id, chunk_index, heading, content) values (%s,%s,%s,%s)",
                            (doc_id, i, item["title"], ch),
                        )
                    conn.commit()
                    existing.add(item["external_id"])
                    new_count += 1
                    print("IMPORTED: " + series + " " + item["title"] + " chunks=" + str(len(chunk_list)) + " is_current=" + str(item["is_current"]))
    print("Import done: " + str(new_count) + " new documents")


def embed_pending(conn, gemini_key):
    embedded = 0
    while time_left() > 30:
        with conn.cursor() as cur:
            cur.execute("select id, content from chunks where embedding is null order by created_at asc limit 20")
            rows = cur.fetchall()
        if not rows:
            break
        for chunk_id, content in rows:
            if time_left() <= 0:
                break
            try:
                vec = embed_text(content, gemini_key)
            except Exception as e:
                print("WARN: embed failed: " + str(chunk_id) + " " + str(e))
                continue
            if not vec:
                continue
            vec_str = "[" + ",".join("%.8f" % x for x in vec) + "]"
            with conn.cursor() as cur2:
                cur2.execute("update chunks set embedding = %s::vector where id = %s", (vec_str, chunk_id))
            conn.commit()
            embedded += 1
    print("Embedding done: " + str(embedded) + " chunks")


def main():
    conn = psycopg2.connect(NEON_DB_URL, connect_timeout=15)
    try:
        ensure_schema(conn)
        import_new_documents(conn)
        gemini_key = get_admin_gemini_key()
        embed_pending(conn, gemini_key)
    finally:
        conn.close()


if __name__ == "__main__":
    main()
