"""
Sync MF metodiky (Centralni harmonizacni jednotka) do samostatne Neon databaze.

Duvod samostatne databaze: metodicke dokumenty od statnich uradu/komor (mimo
e-Sbirku) jsou novy druh zdroje, ktery chceme postupne rozsirovat (Celni
sprava, Financni sprava, komory...) - aby se neopakoval problem s
dochazejicim mistem v hlavni 500MB Supabase databazi. Kazdy novy zdroj
tohoto typu muze dostat vlastni Neon projekt zdarma. Tato databaze sama o
sobe je navrzena obecne (sloupce "source"/"series"), takze do ni muzou
pribyt i dalsi MF serie (ne jen CHJ) bez zmeny schematu.

Na rozdil od e-Sbirka/UOHS/NS pipeline (ktere pisi pres Supabase PostgREST)
tahle jede pres primo Postgres spojeni (Neon nema vestaveny REST API), a
protoze zadna sdilena embedovaci fronta pro tuto databazi neexistuje,
embedding se dela rovnou v tomto skriptu (FAZE A+B dohromady, ne oddelene
jako u ostatnich zdroju).

Zdroj dat: mf.gov.cz/cs/kontrola-a-regulace/rizeni-a-kontrola-verejnych-financi/metodicke-materialy-chj/{rok}
- staticke HTML seznamy po letech, kazda polozka ma detail stranku s jednim
PDF odkazem (overeno rucne 2026-08-07).

MAX_NEW_DOCS_PER_RUN a TIME_BUDGET_SECONDS omezuji beh (stejny vzorec jako u
ostatnich sync skriptu) - skript preskakuje uz existujici zaznamy podle
external_id (slug z URL), takze dalsi beh (mesicni cron) pokracuje tam, kde
skoncil predchozi. Prvni beh projede zpetne az do START_YEAR (zaklad pro
zpetne dohnani historie), dalsi mesicni behy uz jen doplni nove polozky.
"""

import os
import re
import time
from datetime import datetime
from io import BytesIO
from docx import Document

import requests
from bs4 import BeautifulSoup
from pypdf import PdfReader
import psycopg2

NEON_DB_URL = os.environ["NEON_MF_DB_URL"]
SUPABASE_URL = os.environ["SUPABASE_URL"].rstrip("/")
SERVICE_KEY = os.environ["SUPABASE_SERVICE_ROLE_KEY"]
ADMIN_USER_ID = "2648f5db-bea6-4cac-b490-ad0ec59723df"
GEMINI_API_KEY_OVERRIDE = os.environ.get("GEMINI_API_KEY_OVERRIDE")

EMBED_MODEL = "gemini-embedding-001"
EMBED_DIM = 256
MAX_CHARS_PER_CHUNK = 3000
MAX_NEW_DOCS_PER_RUN = int(os.environ.get("MAX_NEW_DOCS_PER_RUN", "40"))
TIME_BUDGET_SECONDS = int(os.environ.get("TIME_BUDGET_SECONDS", "3000"))
START_YEAR = int(os.environ.get("START_YEAR", "2015"))

BASE_URL = "https://mf.gov.cz"
LIST_URL_TMPL = (
    BASE_URL
    + "/cs/kontrola-a-regulace/rizeni-a-kontrola-verejnych-financi/metodicke-materialy-chj/{year}"
)

DETAIL_RE = re.compile(r"/metodicke-materialy-chj/\d{4}/[a-z0-9-]+", re.I)

START_TIME = time.time()


def time_left():
    return TIME_BUDGET_SECONDS - (time.time() - START_TIME)


def sb_headers():
    return {
        "apikey": SERVICE_KEY,
        "Authorization": f"Bearer {SERVICE_KEY}",
        "Content-Type": "application/json",
    }


def get_admin_gemini_key():
    if GEMINI_API_KEY_OVERRIDE:
        return GEMINI_API_KEY_OVERRIDE
    r = requests.post(
        f"{SUPABASE_URL}/rest/v1/rpc/get_user_gemini_key",
        headers=sb_headers(),
        json={"p_user_id": ADMIN_USER_ID},
        timeout=30,
    )
    r.raise_for_status()
    key = r.json()
    if not key:
        raise RuntimeError("Admin Gemini key not available")
    return key


def ensure_schema(conn):
    with conn.cursor() as cur:
        cur.execute("create extension if not exists pgcrypto;")
        cur.execute("create extension if not exists vector;")
        cur.execute(
            """
            create table if not exists documents (
                id uuid primary key default gen_random_uuid(),
                source text not null default 'mf_chj',
                series text,
                external_id text not null,
                title text not null,
                url text not null,
                pdf_url text,
                published_date date,
        is_current boolean not null default true,
                fetched_at timestamptz default now(),
                created_at timestamptz default now(),
                unique(source, external_id)
            );
            """
        )
        cur.execute(
            """
            create table if not exists chunks (
                id uuid primary key default gen_random_uuid(),
                document_id uuid not null references documents(id) on delete cascade,
                chunk_index int not null,
                heading text,
                content text,
                embedding vector(256),
                created_at timestamptz default now()
            );
            """
        )
        cur.execute("create index if not exists chunks_document_id_idx on chunks(document_id);")
        cur.execute(
            "alter table documents add column if not exists is_current boolean not null default true;"
        )
    conn.commit()


def fetch_existing_external_ids(conn, source="mf_chj"):
    with conn.cursor() as cur:
        cur.execute("select external_id from documents where source = %s;", (source,))
        return {row[0] for row in cur.fetchall()}


def slug_from_href(href):
    m = DETAIL_RE.search(href)
    if not m:
        return None
    return m.group(0).split("/")[-1]


def list_year(year):
    resp = requests.get(LIST_URL_TMPL.format(year=year), timeout=60)
    if resp.status_code == 404:
        return []
    resp.raise_for_status()
    soup = BeautifulSoup(resp.text, "html.parser")
    items = []
    seen = set()
    for a in soup.find_all("a", href=True):
        href = a["href"]
        if not DETAIL_RE.search(href):
            continue
        full = href if href.startswith("http") else BASE_URL + href
        slug = slug_from_href(full)
        if not slug or slug in seen:
            continue
        seen.add(slug)
        title = a.get_text(strip=True)
        if not title:
            continue
        items.append({"url": full, "slug": slug, "title": title})
    return items


CZECH_MONTHS = {
    "ledna": 1, "unora": 2, "brezna": 3, "dubna": 4,
    "kvetna": 5, "cervna": 6, "cervence": 7,
    "srpna": 8, "zari": 9, "rijna": 10, "listopadu": 11,
    "prosince": 12,
}


def _strip_diacritics(s):
    table = str.maketrans("ÃÂÃÂÃÂÃÂ¡ÃÂÃÂÃÂÃÂÃÂÃÂÃÂÃÂÃÂÃÂÃÂÃÂ©ÃÂÃÂÃÂÃÂÃÂÃÂÃÂÃÂ­ÃÂÃÂÃÂÃÂÃÂÃÂÃÂÃÂ³ÃÂÃÂÃÂÃÂÃÂÃÂÃÂÃÂ¡ÃÂÃÂÃÂÃÂ¥ÃÂÃÂÃÂÃÂºÃÂÃÂÃÂÃÂ¯ÃÂÃÂÃÂÃÂ½ÃÂÃÂÃÂÃÂ¾", "acdeeinorstuuyz")
    return s.lower().translate(table)


def parse_czech_date(text):
    m = re.search(r"(\d{1,2})\.?\s*(\d{1,2}|[a-zA-ZÃÂÃÂÃÂÃÂ¡-ÃÂÃÂÃÂÃÂ¾ÃÂÃÂÃÂÃÂ-ÃÂÃÂÃÂÃÂ½]+)\.?\s*(\d{4})", text)
    if not m:
        return None
    day, month_raw, year = m.groups()
    month_raw_l = _strip_diacritics(month_raw.strip("."))
    if month_raw_l.isdigit():
        month = int(month_raw_l)
    else:
        month = CZECH_MONTHS.get(month_raw_l)
    if not month:
        return None
    try:
        return datetime(int(year), month, int(day)).date()
    except ValueError:
        return None


def fetch_detail(url, title=None):
    resp = requests.get(url, timeout=60)
    resp.raise_for_status()
    soup = BeautifulSoup(resp.text, "html.parser")
    article = soup.find("article") or soup
    # NOTE: attachment links (PDF and DOCX) live OUTSIDE <article> on
    # mf.gov.cz detail pages (in a .b-download quick-box and a "Dokumenty ke
    # stazeni" archive list), so we must search the whole page, not just the
    # article tag.
    def _collect(ext_regex):
        out = []
        for a in soup.find_all("a", href=True):
            if re.search(ext_regex, a["href"], re.I):
                href = a["href"] if a["href"].startswith("http") else BASE_URL + a["href"]
                out.append((href, a.get_text(" ", strip=True)))
        return out

    def _pick(candidates):
        if not candidates:
            return None
        ident = None
        if title:
            m = re.search(r"ÃÂ\.?\s*(\d+[a-z]?)\s*/\s*(\d{4})", title, re.I)
            if m:
                ident = f"{m.group(1)}/{m.group(2)}"
        if ident:
            for href, text in candidates:
                if ident in text.replace(" ", ""):
                    return href
        # fall back to the last match on the page: observed pattern is
        # superseded/older versions listed first, current version last
        return candidates[-1][0]

    pdf_link = _pick(_collect(r"\.pdf($|\?)"))
    docx_link = _pick(_collect(r"\.docx($|\?)"))
    if pdf_link:
        file_url, file_kind = pdf_link, "pdf"
    elif docx_link:
        file_url, file_kind = docx_link, "docx"
    else:
        file_url, file_kind = None, None
    date_el_text = article.get_text(" ", strip=True)
    published = parse_czech_date(date_el_text)
    # supersession note, e.g. title says "puvodne stanovisko CHJ c. 5/2018"
    replaces_ident = None
    if title:
        m2 = re.search(r"p[uÃÂ¯]vodn[ÃÂe][^0-9]{0,40}ÃÂ\.?\s*(\d+[a-z]?)\s*/\s*(\d{4})", title, re.I)
        if m2:
            replaces_ident = f"{m2.group(1)}/{m2.group(2)}"
    return {
        "pdf_url": file_url,
        "file_kind": file_kind,
        "published_date": published,
        "replaces_ident": replaces_ident,
    }


def extract_pdf_text(pdf_bytes):
    reader = PdfReader(BytesIO(pdf_bytes))
    parts = []
    for page in reader.pages:
        t = page.extract_text() or ""
        parts.append(t)
    text = "\n".join(parts)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_docx_text(docx_bytes):
    doc = Document(BytesIO(docx_bytes))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts).strip()


def split_into_chunks(text, max_chars=MAX_CHARS_PER_CHUNK):
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks = []
    current = []
    current_len = 0
    for p in paragraphs:
        if current_len + len(p) > max_chars and current:
            chunks.append("\n\n".join(current))
            current = []
            current_len = 0
        current.append(p)
        current_len += len(p)
    if current:
        chunks.append("\n\n".join(current))
    return chunks or ([text] if text else [])


def embed_text(text, gemini_key, retries=3):
    for attempt in range(retries):
        try:
            resp = requests.post(
                f"https://generativelanguage.googleapis.com/v1beta/models/{EMBED_MODEL}:embedContent",
                headers={"Content-Type": "application/json", "x-goog-api-key": gemini_key},
                json={
                    "content": {"parts": [{"text": text[:8000]}]},
                    "taskType": "RETRIEVAL_DOCUMENT",
                    "outputDimensionality": EMBED_DIM,
                },
                timeout=30,
            )
            if resp.status_code == 429:
                time.sleep(5 * (attempt + 1))
                continue
            resp.raise_for_status()
            vec = resp.json().get("embedding", {}).get("values")
            if not vec:
                return None
            return vec
        except requests.RequestException:
            if attempt == retries - 1:
                raise
            time.sleep(3 * (attempt + 1))
    return None


def import_new_documents(conn):
    existing = fetch_existing_external_ids(conn)
    current_year = datetime.now().year
    new_count = 0
    for year in range(current_year, START_YEAR - 1, -1):
        if time_left() <= 0 or new_count >= MAX_NEW_DOCS_PER_RUN:
            break
        try:
            items = list_year(year)
        except requests.RequestException as e:
            print(f"WARN: list_year({year}) failed: {e}")
            continue
        for item in items:
            if time_left() <= 0 or new_count >= MAX_NEW_DOCS_PER_RUN:
                break
            if item["slug"] in existing:
                continue
            try:
                detail = fetch_detail(item["url"], title=item.get("title"))
            except requests.RequestException as e:
                print(f"WARN: fetch_detail({item['url']}) failed: {e}")
                continue
            if not detail["pdf_url"]:
                print(f"SKIP (no pdf/docx): {item['title']}")
                existing.add(item["slug"])
                continue
            try:
                file_resp = requests.get(detail["pdf_url"], timeout=60)
                file_resp.raise_for_status()
                if detail.get("file_kind") == "docx":
                    text = extract_docx_text(file_resp.content)
                else:
                    text = extract_pdf_text(file_resp.content)
            except Exception as e:
                print(f"WARN: file fetch/extract failed for {item['title']}: {e}")
                continue
            if not text:
                print(f"SKIP (empty text): {item['title']}")
                existing.add(item["slug"])
                continue
            chunk_texts = split_into_chunks(text)
            with conn.cursor() as cur:
                cur.execute(
                    """
                    insert into documents (source, series, external_id, title, url, pdf_url, published_date)
                    values (%s, %s, %s, %s, %s, %s, %s)
                    on conflict (source, external_id) do nothing
                    returning id;
                    """,
                    (
                        "mf_chj",
                        "chj_metodika",
                        item["slug"],
                        item["title"],
                        item["url"],
                        detail["pdf_url"],
                        detail["published_date"],
                    ),
                )
                row = cur.fetchone()
                if not row:
                    conn.commit()
                    continue
                doc_id = row[0]
                if detail.get("replaces_ident"):
                    print(f"SUPERSEDES: {item['title']} -> marking c. {detail['replaces_ident']} as historical")
                    old_num, old_year = detail["replaces_ident"].split("/")
                    pat = r"ÃÂ\.?\s*" + re.escape(old_num) + r"\s*/\s*" + re.escape(old_year)
                    cur.execute(
                        """
                        update documents set is_current = false
                        where source = %s and is_current = true and id != %s
                          and title ~* %s
                        """,
                        ("mf_chj", doc_id, pat),
                    )
                for i, ctext in enumerate(chunk_texts):
                    heading = item["title"] if i == 0 else f"{item['title']} (pokracovani {i + 1})"
                    cur.execute(
                        """
                        insert into chunks (document_id, chunk_index, heading, content)
                        values (%s, %s, %s, %s);
                        """,
                        (doc_id, i, heading, ctext),
                    )
            conn.commit()
            existing.add(item["slug"])
            new_count += 1
            print(f"IMPORTED: {item['title']} ({len(chunk_texts)} chunks)")
    print(f"Import done: {new_count} new documents")
    return new_count


def embed_pending(conn, gemini_key):
    embedded = 0
    while time_left() > 30:
        with conn.cursor() as cur:
            cur.execute(
                "select id, content from chunks where embedding is null order by created_at asc limit 20;"
            )
            rows = cur.fetchall()
        if not rows:
            break
        for chunk_id, content in rows:
            if time_left() <= 0:
                break
            try:
                vec = embed_text(content, gemini_key)
            except Exception as e:
                print(f"WARN: embed failed for {chunk_id}: {e}")
                continue
            if not vec:
                continue
            vec_str = "[" + ",".join(f"{x:.8f}" for x in vec) + "]"
            with conn.cursor() as cur:
                cur.execute(
                    "update chunks set embedding = %s::vector where id = %s;",
                    (vec_str, chunk_id),
                )
            conn.commit()
            embedded += 1
    print(f"Embedding done: {embedded} chunks")
    return embedded


def main():
    conn = psycopg2.connect(NEON_DB_URL, connect_timeout=15)
    try:
        ensure_schema(conn)
        import_new_documents(conn)
        gemini_key = get_admin_gemini_key()
        embed_pending(conn, gemini_key)
    finally:
        conn.close()


if __name__ == "__main__":
    main()
